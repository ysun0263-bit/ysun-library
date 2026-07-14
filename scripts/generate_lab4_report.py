from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE = ROOT / "evidence"
DOCX = EVIDENCE / "FIT5032_Assessed_Lab_4_Yilian_Sun.docx"
PDF = EVIDENCE / "FIT5032_Assessed_Lab_4_Yilian_Sun.pdf"


FIGURES = [
    ("01_activity1_original_form.png", "Figure 1. Original Week 3 form allows an empty submission and creates an empty card."),
    ("02_activity2_required_username.png", "Figure 2. HTML built-in validation prevents submission when Username is empty."),
    ("03_activity2_short_password.png", "Figure 3. HTML built-in validation evidence for password length checking with password value 123."),
    ("04_activity3_username_error.png", "Figure 4. Vue validation prevents a missing username from being submitted."),
    ("05_activity3_password_error.png", "Figure 5. Vue password validation reports an invalid password and blocks submission."),
    ("06_activity3_extra_validations.png", "Figure 6. Vue validation displays additional Resident, Gender, and Reason errors."),
    ("07_valid_submission.png", "Figure 7. A valid record is accepted and displayed after passing all five validations."),
    ("08_primevue_datatable.png", "Figure 8. PrimeVue DataTable displays three valid submitted user records."),
    ("09_validation_code.png", "Figure 9. Source evidence for the five Vue validation functions."),
    ("10_datatable_code.png", "Figure 10. Source evidence for PrimeVue DataTable imports and template."),
    ("11_primevue_config_code.png", "Figure 11. Source evidence for PrimeVue Aura configuration in main.js."),
    ("12_github_commit_history.png", "Figure 12. Git commit history showing staged assessed lab updates."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FIT5032 Assessed Lab 4")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("HTML/Vue Form Validation and PrimeVue DataTable").italic = True

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Student", "Yilian Sun"),
        ("Student ID", "[Student ID to be completed]"),
        ("Project", "FIT5032 Week 3 Library Vue application"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph(
        "This report records the implemented assessed lab work, source evidence, testing outcomes, and Git history for the Library Vue project."
    )
    doc.add_page_break()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_figure(doc, filename, caption):
    image_path = EVIDENCE / filename
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(6.2))
    para = doc.add_paragraph(caption)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].italic = True


def create_commit_history_image():
    log = subprocess.check_output(
        ["git", "log", "--oneline", "--decorate", "-n", "10"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
    )
    font_path = Path("C:/Windows/Fonts/consola.ttf")
    font = ImageFont.truetype(str(font_path), 22) if font_path.exists() else ImageFont.load_default()
    title_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 26)
    lines = ["Git commit history", ""] + log.strip().splitlines()
    width = 1500
    height = 90 + len(lines) * 34
    image = Image.new("RGB", (width, height), "#eef2f7")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((34, 34, width - 34, height - 34), radius=10, fill="#101827", outline="#263244", width=2)
    draw.text((60, 52), lines[0], font=title_font, fill="#f8fafc")
    y = 104
    for line in lines[2:]:
        draw.text((60, y), line, font=font, fill="#dbeafe")
        y += 34
    image.save(EVIDENCE / "12_github_commit_history.png")


def build_report():
    create_commit_history_image()
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("eFolio Task 4.1", level=1)
    doc.add_paragraph(
        "Activities 1-3 were completed in sequence: the original form was captured, HTML built-in validation was added, and the final implementation was converted to Vue custom validation."
    )
    for filename, caption in FIGURES[:6]:
        add_figure(doc, filename, caption)

    doc.add_heading("eFolio Task 4.2", level=1)
    doc.add_paragraph(
        "The final form implements five field-level validations and displays accepted records in a PrimeVue DataTable. Invalid data is blocked before it can enter the Card list or DataTable."
    )
    add_table(
        doc,
        ["Field", "Final validation rule"],
        [
            ("Username", "Trimmed; required; at least 3 characters; spaces-only input rejected."),
            ("Password", "At least 8 characters; includes uppercase, lowercase, number, and special character."),
            ("Australian Resident", "User must explicitly choose Yes or No from a blank default select."),
            ("Gender", "User must explicitly choose Male, Female, or Other from a blank default select."),
            ("Reason", "Trimmed; required; 10 to 200 characters."),
        ],
    )
    for filename, caption in FIGURES[6:12]:
        add_figure(doc, filename, caption)

    doc.add_heading("Testing Summary", level=1)
    add_table(
        doc,
        ["Scenario", "Result", "Evidence"],
        [
            ("Original empty submission creates an empty card before validation", "Pass", "Figure 1"),
            ("HTML required username blocks empty submit", "Pass", "Figure 2"),
            ("HTML password length rule added and checked with value 123", "Pass", "Figure 3"),
            ("Vue empty/short username blocks submit", "Pass", "Figure 4"),
            ("Vue password complexity blocks invalid passwords", "Pass", "Figure 5"),
            ("Resident, Gender, and Reason extra validations block submit", "Pass", "Figure 6"),
            ("Valid data submits successfully", "Pass", "Figure 7"),
            ("Three valid submissions display as three DataTable rows", "Pass", "Figure 8"),
            ("Browser console errors during DataTable capture", "Pass - none observed", "Browser automation log"),
            ("npm run build", "Pass", "Vite production build completed successfully"),
        ],
    )

    doc.add_heading("Build and Git Summary", level=1)
    doc.add_paragraph("Build result: npm run build completed successfully after PrimeVue integration.")
    doc.add_paragraph(
        "Git commits were created for HTML validation, Vue validation, and PrimeVue DataTable work. A final documentation commit will include this report and evidence."
    )
    doc.add_paragraph(
        "Student action still required: fill in the Student ID placeholder, confirm tutor sharing on the private GitHub repository, and upload the final PDF to Moodle."
    )

    doc.save(DOCX)
    build_pdf()


def scaled_pdf_image(path, max_width=6.5 * inch, max_height=7.2 * inch):
    with Image.open(path) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return PdfImage(str(path), width=width * scale, height=height * scale)


def pdf_table(data):
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            alignment=1,
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#374151"),
            spaceAfter=10,
        )
    )
    styles["Title"].textColor = colors.HexColor("#0B2545")
    styles["Heading1"].textColor = colors.HexColor("#2E74B5")
    styles["Heading2"].textColor = colors.HexColor("#2E74B5")

    story = []
    story.append(Paragraph("FIT5032 Assessed Lab 4", styles["Title"]))
    story.append(Paragraph("Yilian Sun", styles["Heading2"]))
    story.append(Paragraph("Student ID: [Student ID to be completed]", styles["BodyText"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        Paragraph(
            "This PDF contains the real implementation evidence, test outcomes, and Git history for the Week 3 Library Vue project extended for Assessed Lab 4.",
            styles["BodyText"],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("eFolio Task 4.1", styles["Heading1"]))
    story.append(
        Paragraph(
            "The original form, HTML built-in validation stage, and final Vue custom validation stage were implemented and captured below.",
            styles["BodyText"],
        )
    )
    for filename, caption in FIGURES[:6]:
        story.append(scaled_pdf_image(EVIDENCE / filename))
        story.append(Paragraph(caption, styles["Caption"]))

    story.append(PageBreak())
    story.append(Paragraph("eFolio Task 4.2", styles["Heading1"]))
    story.append(
        pdf_table(
            [
                ["Field", "Final validation rule"],
                ["Username", "Trimmed; required; at least 3 characters; spaces-only input rejected."],
                ["Password", "At least 8 characters; includes uppercase, lowercase, number, and special character."],
                ["Australian Resident", "User must explicitly choose Yes or No from a blank default select."],
                ["Gender", "User must explicitly choose Male, Female, or Other from a blank default select."],
                ["Reason", "Trimmed; required; 10 to 200 characters."],
            ]
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    for filename, caption in FIGURES[6:12]:
        story.append(scaled_pdf_image(EVIDENCE / filename))
        story.append(Paragraph(caption, styles["Caption"]))

    story.append(PageBreak())
    story.append(Paragraph("Testing Summary", styles["Heading1"]))
    story.append(
        pdf_table(
            [
                ["Scenario", "Result", "Evidence"],
                ["Original empty submission creates an empty card before validation", "Pass", "Figure 1"],
                ["HTML required username blocks empty submit", "Pass", "Figure 2"],
                ["HTML password length rule added and checked with value 123", "Pass", "Figure 3"],
                ["Vue empty/short username blocks submit", "Pass", "Figure 4"],
                ["Vue password complexity blocks invalid passwords", "Pass", "Figure 5"],
                ["Resident, Gender, and Reason extra validations block submit", "Pass", "Figure 6"],
                ["Valid data submits successfully", "Pass", "Figure 7"],
                ["Three valid submissions display as three DataTable rows", "Pass", "Figure 8"],
                ["Browser console errors during DataTable capture", "Pass - none observed", "Browser automation log"],
                ["npm run build", "Pass", "Vite production build completed successfully"],
            ]
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Build and Git Summary", styles["Heading1"]))
    story.append(Paragraph("Build result: npm run build completed successfully after PrimeVue integration.", styles["BodyText"]))
    story.append(
        Paragraph(
            "Student action still required: fill in the Student ID placeholder, confirm tutor sharing on the private GitHub repository, and upload this PDF to Moodle.",
            styles["BodyText"],
        )
    )

    document = SimpleDocTemplate(
        str(PDF),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    document.build(story)


if __name__ == "__main__":
    build_report()
